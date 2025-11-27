from docx import Document
import xlrd
from line_dictionary import Article

import re

pattern_for_index = r'^[А-Яа-яЁё́]+$' #включая знак ударения
pattern_for_sort = r'^[А-Яа-яЁё]+$'
sort_letters = "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"

def is_russian_letter(text):
    return re.match(pattern_for_index, text) is not None

def sort_rus_key(article):
    text = article.rus.lower()
    word = []
    for letter in text:
        if re.match(pattern_for_sort, letter) is not None:
            word += [sort_letters.index(letter)]
    return word

def upper_index(rus):
    ind = 0
    while ind < len(rus) and is_russian_letter(rus[ind]):
        ind += 1
    letter = rus[ind:ind + 1]
    if letter.isdigit():
        prefix = rus[:ind]
        sufix = rus[ind + 1:]
        p.add_run(prefix).font.bold = True
        temp = p.add_run(letter)
        temp.font.superscript = True
        temp.font.bold = True
        p.add_run(sufix + " ").font.bold = True
        return
    p.add_run(rus + " ").font.bold = True
def bold_polisemy(mult):
    if mult[0].isdigit():
        prefix = mult[:2]
        sufix = mult[2:]
        p.add_run(prefix).font.bold = True
        p.add_run(sufix + " ").font.italic = True
        return
def bold_plural(form):
    p.add_run(" ")
    if form:
        pattern = "бир."
        ind = form.find(pattern)
        if ind > -1:
            prefix = form[:ind]
            sufix = form[ind + len(pattern):]
            p.add_run(" " + prefix).font.italic = True
            temp = p.add_run(pattern)
            temp.font.bold = True
            temp.font.italic = True
            p.add_run(sufix).font.italic = True
            return
        p.add_run(form).font.italic = True

def copy_style(cell):
    # cell_text =
    # for char_index in range(1, len(cell_text) + 1):
    #     char = cell.Characters(char_index, 1)
    #     if char.Font.Bold:
    #         range_obj.Characters(char_index).Font.Bold = True
    #     if char.Font.Italic:
    #         range_obj.Characters(char_index).Font.Italic = True
    #     if char.Font.Size:
    #         range_obj.Characters(char_index).Font.Size = char.Font.Size
    #     if char.Font.Name:
    #         range_obj.Characters(char_index).Font.Name = char.Font.Name
    return

workbook = xlrd.open_workbook('+ОСЕТ-РУС+_stress_sort.xlsx')
sheet = workbook.sheet_by_name('Sheet')

document = Document()

# lines = line_dictionary(sheet)
row = 2 #2
# line1 = lines.read_line(1)
# line2 = lines.read_line(2)
# print(line1.ose, line2.ose)
total_row = 790 # sheet.nrows
article = ""
# считываем весь словарь
articles = list()
while row < total_row:
    article = Article(sheet)
    print(row)
    row, art = article.read_article(row)
    print(row, art)
    articles.append(art)
    row += 1
# сортируем статьи

articles.sort(key=sort_rus_key)
# записываем статьи в нужном виде
for ind_dict, art in enumerate(articles):
    print(ind_dict, art.rus)
    p = document.add_paragraph()
    upper_index(art.rus)
    if art.reduct: p.add_run(art.reduct + " ").font.italic = True

    for ind_pol, polisemy in enumerate(art.polisemy):
        if polisemy.multy:
            if polisemy.transc[0][3]: p.add_run(polisemy.transc[0][3] + " ").font.italic = True
            bold_polisemy(polisemy.multy)
            #p.add_run(polisemy.multy + " ")
        for ind, (ose, tr, form, notice) in enumerate(polisemy.transc):
            if not polisemy.multy: p.add_run(notice + " ").font.italic = True
            p.add_run(ose)
            #p.add_run("/" + tr + "/").font.name = 'Ipa-samd Uclphon1 SILDoulosL'
            if tr:
                p.add_run(" /" + tr + "/")
            if form:
                bold_plural(form)
            if ind != len(polisemy.transc) - 1:
                p.add_run(", ")
        for ind, (exm_rus, exm_ose) in enumerate(polisemy.examples):
            p.add_run("; ")
            if exm_rus:
                p.add_run(" " + exm_rus).font.bold = True
                p.add_run(" " + exm_ose)

document.save('test3.docx')